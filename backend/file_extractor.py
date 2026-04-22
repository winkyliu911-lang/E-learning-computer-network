from __future__ import annotations

import os
import uuid
from pathlib import Path
from typing import List

from PIL import Image, ImageDraw, ImageFont


def _pick_cjk_font_path() -> str | None:
	"""优先选择系统中可用的中文字体，避免默认字体导致中文乱码。"""
	candidates = [
		"/System/Library/Fonts/PingFang.ttc",
		"/System/Library/Fonts/Hiragino Sans GB.ttc",
		"/System/Library/Fonts/Supplemental/Songti.ttc",
		"/System/Library/Fonts/STHeiti Light.ttc",
		"/Library/Fonts/Arial Unicode.ttf",
		"/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
		"/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
	]
	for font_path in candidates:
		if os.path.exists(font_path):
			return font_path
	return None


def _load_font(font_size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
	font_path = _pick_cjk_font_path()
	if font_path:
		try:
			return ImageFont.truetype(font_path, font_size)
		except Exception:
			pass
	return ImageFont.load_default()


def _split_paragraph_to_lines(paragraph: str, draw: ImageDraw.ImageDraw, font: ImageFont.ImageFont, max_width: int) -> List[str]:
	if not paragraph:
		return [""]

	lines: List[str] = []
	current = ""
	for ch in paragraph:
		candidate = current + ch
		w = draw.textlength(candidate, font=font)
		if w <= max_width:
			current = candidate
		else:
			if current:
				lines.append(current)
			current = ch
	if current:
		lines.append(current)
	return lines


def _render_text_to_paged_images(text: str, output_dir: str, file_prefix: str) -> List[str]:
	"""将长文本按页渲染为图片，返回图片路径列表。"""
	output_path = Path(output_dir)
	output_path.mkdir(parents=True, exist_ok=True)

	page_width, page_height = 1240, 1754
	margin_x, margin_y = 72, 72
	line_spacing = 10
	font = _load_font(30)

	measure_canvas = Image.new("RGB", (page_width, page_height), "white")
	draw = ImageDraw.Draw(measure_canvas)
	max_line_width = page_width - margin_x * 2

	bbox = draw.textbbox((0, 0), "中A", font=font)
	line_height = (bbox[3] - bbox[1]) + line_spacing
	max_lines_per_page = max(1, (page_height - margin_y * 2) // line_height)

	paragraphs = text.splitlines() if text else [""]
	all_lines: List[str] = []
	for paragraph in paragraphs:
		wrapped = _split_paragraph_to_lines(paragraph, draw, font, max_line_width)
		all_lines.extend(wrapped if wrapped else [""])

	image_paths: List[str] = []
	page_no = 1
	for i in range(0, len(all_lines), max_lines_per_page):
		page_lines = all_lines[i:i + max_lines_per_page]
		canvas = Image.new("RGB", (page_width, page_height), "white")
		page_draw = ImageDraw.Draw(canvas)

		y = margin_y
		for line in page_lines:
			page_draw.text((margin_x, y), line, fill="black", font=font)
			y += line_height

		filename = f"{file_prefix}_page_{page_no:03d}.png"
		full_path = output_path / filename
		canvas.save(full_path, format="PNG")
		image_paths.append(str(full_path))
		page_no += 1

	return image_paths


def convert_pdf_to_images(pdf_path: str, output_dir: str, file_prefix: str) -> List[str]:
	"""将 PDF 每页渲染成 PNG 图片。"""
	try:
		import fitz  # PyMuPDF
	except Exception as e:
		raise RuntimeError("缺少 PyMuPDF 依赖，请安装: pip install pymupdf") from e

	output_path = Path(output_dir)
	output_path.mkdir(parents=True, exist_ok=True)

	image_paths: List[str] = []
	with fitz.open(pdf_path) as doc:
		for page_idx in range(len(doc)):
			page = doc[page_idx]
			pix = page.get_pixmap(matrix=fitz.Matrix(1.8, 1.8), alpha=False)
			filename = f"{file_prefix}_page_{page_idx + 1:03d}.png"
			full_path = output_path / filename
			pix.save(str(full_path))
			image_paths.append(str(full_path))

	return image_paths


def convert_docx_to_images(docx_path: str, output_dir: str, file_prefix: str) -> List[str]:
	"""将 DOCX 文本渲染为分页图片。"""
	try:
		from docx import Document
	except Exception as e:
		raise RuntimeError("缺少 python-docx 依赖，请安装: pip install python-docx") from e

	doc = Document(docx_path)

	blocks: List[str] = []
	for p in doc.paragraphs:
		txt = (p.text or "").strip()
		if txt:
			blocks.append(txt)

	for table in doc.tables:
		for row in table.rows:
			row_text = " | ".join((cell.text or "").strip() for cell in row.cells)
			if row_text.strip(" |"):
				blocks.append(row_text)

	merged_text = "\n".join(blocks).strip()
	if not merged_text:
		merged_text = "（该 Word 文件未提取到可显示文本）"

	return _render_text_to_paged_images(merged_text, output_dir, file_prefix)


def convert_upload_to_images(file_path: str, output_dir: str, owner_prefix: str) -> List[str]:
	"""统一入口：将上传文件转换为一组图片路径。"""
	src = Path(file_path)
	ext = src.suffix.lower()
	unique_prefix = f"{owner_prefix}_{uuid.uuid4().hex[:8]}_{src.stem}"

	if ext in {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"}:
		return [str(src)]
	if ext == ".pdf":
		return convert_pdf_to_images(str(src), output_dir, unique_prefix)
	if ext == ".docx":
		return convert_docx_to_images(str(src), output_dir, unique_prefix)

	raise ValueError(f"不支持的文件格式: {src.name}")

