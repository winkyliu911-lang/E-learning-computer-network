"""
RAG Manager - 管理 RAG 系统的核心类
处理文档加载、向量化和检索
支持混合检索：BM25 关键词搜索 + 向量语义搜索 + RRF 融合
"""

from __future__ import annotations

import logging
import math
import pickle
import re
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional

import chromadb

logger = logging.getLogger(__name__)

# LangChain / 依赖兼容导入

try:
    from langchain_core.documents import Document
except Exception:
    try:
        from langchain.schema import Document
    except Exception:
        class Document:
            def __init__(self, page_content: str, metadata: dict = None):
                self.page_content = page_content
                self.metadata = metadata or {}

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except Exception:
    try:
        from langchain.text_splitter import RecursiveCharacterTextSplitter
    except Exception:
        RecursiveCharacterTextSplitter = None

try:
    from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, TextLoader
except Exception:
    try:
        from langchain.document_loaders import PyPDFLoader, Docx2txtLoader, TextLoader
    except Exception:
        PyPDFLoader = Docx2txtLoader = TextLoader = None

try:
    from langchain_chroma import Chroma
except Exception:
    try:
        from langchain.vectorstores import Chroma
    except Exception:
        Chroma = None

HuggingFaceEmbeddings = None
SentenceTransformer = None
try:
    from langchain_huggingface import HuggingFaceEmbeddings as _HFE
    HuggingFaceEmbeddings = _HFE
except Exception:
    try:
        from langchain.embeddings import HuggingFaceEmbeddings as _HFE
        HuggingFaceEmbeddings = _HFE
    except Exception:
        try:
            from sentence_transformers import SentenceTransformer
        except Exception:
            SentenceTransformer = None

# BM25 + 中文分词
try:
    from rank_bm25 import BM25Okapi
except ImportError:
    BM25Okapi = None

try:
    import jieba
except ImportError:
    jieba = None


# =========================
# 常量与辅助函数
# =========================

RRF_K = 60

_STOPWORDS_ZH = {
    "的", "了", "是", "在", "我", "有", "和", "就", "不", "人", "都", "一", "一个",
    "上", "也", "很", "到", "说", "要", "去", "你", "会", "着", "没有", "看", "好",
    "自己", "这", "那", "被", "把", "从", "对", "他", "她", "它", "们", "什么",
    "如何", "为什么", "怎么", "哪", "吗", "呢", "吧", "啊", "嗯",
}

_QUERY_EXPANSIONS = {
    "三次握手": "TCP SYN SYN-ACK ACK three-way handshake",
    "四次挥手": "TCP FIN ACK four-way handshake",
    "域名": "DNS domain name",
    "路由": "routing OSPF RIP BGP",
    "子网掩码": "subnet mask CIDR",
    "滑动窗口": "sliding window TCP flow control",
    "拥塞控制": "congestion control TCP",
    "物理层": "physical layer",
    "数据链路层": "data link layer MAC",
    "网络层": "network layer IP",
    "传输层": "transport layer TCP UDP",
    "应用层": "application layer HTTP DNS SMTP",
    "交换机": "switch MAC",
    "路由器": "router",
    "防火墙": "firewall",
    "NAT": "network address translation",
    "ARP": "address resolution protocol",
    "ICMP": "internet control message protocol ping",
    "HTTP": "hypertext transfer protocol",
    "HTTPS": "SSL TLS",
}


def _tokenize(text: str) -> List[str]:
    if not text:
        return []
    if jieba is not None:
        tokens = list(jieba.cut_for_search(text))
    else:
        tokens = text.split()
    result = []
    for t in tokens:
        t = t.strip().lower()
        if len(t) < 1:
            continue
        if t in _STOPWORDS_ZH:
            continue
        if re.fullmatch(r'[\s\W]+', t):
            continue
        result.append(t)
    return result


def _expand_query(query: str) -> str:
    expanded = query
    for zh_term, expansion in _QUERY_EXPANSIONS.items():
        if zh_term in query:
            expanded += " " + expansion
    return expanded


# =========================
# Embeddings 统一封装
# =========================

class _STFallbackEmbeddings:
    def __init__(self, model_name: str):
        if SentenceTransformer is None:
            raise ImportError("缺少 embeddings 依赖。请安装：pip install -U sentence-transformers")
        self.model = SentenceTransformer(model_name)

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        embs = self.model.encode(texts, convert_to_numpy=True)
        return [e.tolist() if hasattr(e, "tolist") else list(e) for e in embs]

    def embed_query(self, text: str) -> List[float]:
        emb = self.model.encode([text], convert_to_numpy=True)[0]
        return emb.tolist() if hasattr(emb, "tolist") else list(emb)


def _build_embeddings(model_name: str):
    if HuggingFaceEmbeddings is not None:
        try:
            return HuggingFaceEmbeddings(model_name=model_name)
        except Exception as e:
            logger.warning(f"HuggingFaceEmbeddings 初始化失败，回退 sentence-transformers: {e}")
    return _STFallbackEmbeddings(model_name=model_name)


# =========================
# RAGManager
# =========================

class RAGManager:

    def __init__(self, vector_db_path: str = "./vector_db"):
        self.vector_db_path = vector_db_path

        try:
            chromadb.configure(anonymized_telemetry=False)
        except Exception:
            pass

        # ChromaDB 0.3.x 使用 Settings 持久化
        try:
            from chromadb.config import Settings
            self.chroma_client = chromadb.Client(Settings(
                chroma_db_impl="duckdb+parquet",
                persist_directory=vector_db_path,
                anonymized_telemetry=False,
            ))
            logger.info(f"ChromaDB 持久化客户端初始化成功: {vector_db_path}")
        except Exception:
            try:
                self.chroma_client = chromadb.PersistentClient(path=vector_db_path)
                logger.info("ChromaDB PersistentClient 初始化成功")
            except Exception:
                try:
                    self.chroma_client = chromadb.Client()
                    logger.warning("使用默认 ChromaDB Client（非持久化）")
                except Exception as e:
                    logger.error(f"初始化 Chroma 客户端失败: {e}")
                    raise

        self.embeddings = _build_embeddings(
            model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
        )

        self.collection_name = "documents"
        self.collection = self.chroma_client.get_or_create_collection(
            name=self.collection_name,
            metadata={"hnsw:space": "cosine"},
        )

        self.vectorstore = self._init_vectorstore()

        # BM25 索引
        self._bm25_index: Optional[Any] = None
        self._bm25_texts: List[str] = []
        self._bm25_metadatas: List[dict] = []
        self._bm25_pickle_path = str(Path(vector_db_path) / "bm25_index.pkl")
        self._load_or_build_bm25()

        logger.info(f"RAG Manager 初始化完成，向量库路径: {vector_db_path}")

    def _init_vectorstore(self):
        """使用 LocalChroma 避免 LangChain Chroma 调用 upsert 的兼容问题"""

        class LocalChroma:
            def __init__(self, collection, embeddings):
                self.collection = collection
                self.embeddings = embeddings

            def add_documents(self, docs: List[Document]):
                ids, documents, metadatas = [], [], []
                for d in docs:
                    ids.append(str(uuid.uuid4()))
                    documents.append(d.page_content if hasattr(d, "page_content") else str(d))
                    metadatas.append(d.metadata if hasattr(d, "metadata") else {})
                embeddings = self.embeddings.embed_documents(documents)
                return self.collection.add(ids=ids, documents=documents, metadatas=metadatas, embeddings=embeddings)

            def similarity_search_with_score(self, query: str, k: int = 5):
                q_emb = self.embeddings.embed_query(query)
                res = self.collection.query(query_embeddings=[q_emb], n_results=k, include=["documents", "metadatas", "distances"])
                docs_list = (res.get("documents") or [[]])[0]
                metas_list = (res.get("metadatas") or [[]])[0]
                dists_list = (res.get("distances") or [[]])[0]
                out = []
                for doc_text, meta, dist in zip(docs_list, metas_list, dists_list):
                    out.append((Document(page_content=doc_text, metadata=meta), dist))
                return out

        logger.info("向量存储初始化完成（LocalChroma）")
        return LocalChroma(self.collection, self.embeddings)

    def _persist(self):
        try:
            if hasattr(self.chroma_client, 'persist'):
                self.chroma_client.persist()
        except Exception:
            pass

    # -------------------------
    # BM25 索引管理
    # -------------------------

    def _load_or_build_bm25(self):
        if BM25Okapi is None:
            logger.warning("rank_bm25 未安装，BM25 搜索不可用")
            return

        try:
            pkl_path = Path(self._bm25_pickle_path)
            if pkl_path.exists():
                with open(pkl_path, "rb") as f:
                    data = pickle.load(f)
                self._bm25_texts = data["texts"]
                self._bm25_metadatas = data["metadatas"]
                corpus_tokens = data["corpus_tokens"]
                if corpus_tokens:
                    self._bm25_index = BM25Okapi(corpus_tokens)
                    logger.info(f"BM25 索引从缓存加载（{len(self._bm25_texts)} 个文本块）")
                    return
        except Exception as e:
            logger.warning(f"BM25 缓存加载失败，重建索引: {e}")

        self._rebuild_bm25_from_chroma()

    def _rebuild_bm25_from_chroma(self):
        if BM25Okapi is None:
            return
        try:
            count = self.collection.count()
            if count == 0:
                logger.info("ChromaDB 为空，跳过 BM25 索引构建")
                return

            all_texts, all_metadatas = [], []
            batch_size = 5000
            offset = 0
            while offset < count:
                batch = self.collection.get(limit=batch_size, offset=offset, include=["documents", "metadatas"])
                docs = batch.get("documents") or []
                metas = batch.get("metadatas") or []
                all_texts.extend(docs)
                all_metadatas.extend(metas)
                offset += len(docs)
                if len(docs) < batch_size:
                    break

            self._bm25_texts = all_texts
            self._bm25_metadatas = all_metadatas
            corpus_tokens = [_tokenize(t) for t in all_texts]
            self._bm25_index = BM25Okapi(corpus_tokens)
            self._save_bm25_pickle(corpus_tokens)
            logger.info(f"BM25 索引已从 ChromaDB 重建（{len(all_texts)} 个文本块）")
        except Exception as e:
            logger.error(f"BM25 索引重建失败: {e}")

    def _update_bm25_with_chunks(self, chunks: List):
        if BM25Okapi is None:
            return
        new_texts = [c.page_content if hasattr(c, "page_content") else str(c) for c in chunks]
        new_metas = [c.metadata if hasattr(c, "metadata") else {} for c in chunks]
        self._bm25_texts.extend(new_texts)
        self._bm25_metadatas.extend(new_metas)
        corpus_tokens = [_tokenize(t) for t in self._bm25_texts]
        if corpus_tokens:
            self._bm25_index = BM25Okapi(corpus_tokens)
            self._save_bm25_pickle(corpus_tokens)
            logger.info(f"BM25 索引已更新（共 {len(self._bm25_texts)} 个文本块）")

    def _save_bm25_pickle(self, corpus_tokens: List[List[str]]):
        try:
            data = {"texts": self._bm25_texts, "metadatas": self._bm25_metadatas, "corpus_tokens": corpus_tokens}
            with open(self._bm25_pickle_path, "wb") as f:
                pickle.dump(data, f)
        except Exception as e:
            logger.warning(f"BM25 缓存保存失败: {e}")

    # -------------------------
    # 搜索方法
    # -------------------------

    def _bm25_search(self, query: str, k: int = 5) -> List[tuple]:
        if self._bm25_index is None or not self._bm25_texts:
            return []
        expanded = _expand_query(query)
        tokens = _tokenize(expanded)
        if not tokens:
            return []
        scores = self._bm25_index.get_scores(tokens)
        top_indices = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)[:k]
        results = []
        for idx in top_indices:
            if scores[idx] <= 0:
                continue
            doc = Document(
                page_content=self._bm25_texts[idx],
                metadata=self._bm25_metadatas[idx] if idx < len(self._bm25_metadatas) else {},
            )
            results.append((doc, float(scores[idx])))
        return results

    def _hybrid_search(self, query: str, k: int = 5) -> List[tuple]:
        n_candidates = k * 3

        try:
            vector_results = self.vectorstore.similarity_search_with_score(query, k=n_candidates)
        except Exception:
            vector_results = []

        bm25_results = self._bm25_search(query, k=n_candidates)

        if not vector_results and not bm25_results:
            return []
        if not bm25_results:
            return vector_results[:k]
        if not vector_results:
            return bm25_results[:k]

        doc_scores: Dict[str, Dict] = {}

        for rank, (doc, _dist) in enumerate(vector_results):
            key = doc.page_content[:200]
            if key not in doc_scores:
                doc_scores[key] = {"doc": doc, "rrf_score": 0.0}
            doc_scores[key]["rrf_score"] += 1.0 / (RRF_K + rank + 1)

        for rank, (doc, _bm25_score) in enumerate(bm25_results):
            key = doc.page_content[:200]
            if key not in doc_scores:
                doc_scores[key] = {"doc": doc, "rrf_score": 0.0}
            doc_scores[key]["rrf_score"] += 1.0 / (RRF_K + rank + 1)

        sorted_docs = sorted(doc_scores.values(), key=lambda x: x["rrf_score"], reverse=True)[:k]
        return [(entry["doc"], entry["rrf_score"]) for entry in sorted_docs]

    # -------------------------
    # Public APIs
    # -------------------------

    def add_documents(self, file_paths: List[str], document_source: str = "user") -> Dict[str, Any]:
        if RecursiveCharacterTextSplitter is None:
            raise ImportError("RecursiveCharacterTextSplitter 未成功导入。请安装：pip install -U langchain-text-splitters")

        all_docs: List[Document] = []
        errors: List[str] = []
        processed_files = 0

        for fp in file_paths:
            file_path = Path(fp)
            try:
                suf = file_path.suffix.lower()
                if suf == ".pdf":
                    if PyPDFLoader is None:
                        raise ImportError("PyPDFLoader 未导入。请安装：pip install -U langchain-community pypdf")
                    logger.info(f"加载 PDF: {file_path.name}")
                    loader = PyPDFLoader(str(file_path))
                    docs = loader.load()
                    for d in docs:
                        md = getattr(d, "metadata", {}) or {}
                        md.update({"source": document_source, "file": file_path.name})
                        d.metadata = md
                    all_docs.extend(docs)
                    processed_files += 1

                elif suf == ".docx":
                    logger.info(f"加载 Word: {file_path.name}")
                    if Docx2txtLoader is not None:
                        loader = Docx2txtLoader(str(file_path))
                        docs = loader.load()
                        for d in docs:
                            md = getattr(d, "metadata", {}) or {}
                            md.update({"source": document_source, "file": file_path.name})
                            d.metadata = md
                        all_docs.extend(docs)
                    else:
                        try:
                            from docx import Document as _Docx
                            docx_obj = _Docx(str(file_path))
                            full = "\n".join([p.text for p in docx_obj.paragraphs])
                            all_docs.append(Document(page_content=full, metadata={"source": document_source, "file": file_path.name}))
                        except Exception as ex:
                            raise RuntimeError("读取 .docx 失败。请安装：pip install -U python-docx") from ex
                    processed_files += 1

                elif suf == ".md":
                    logger.info(f"加载 Markdown: {file_path.name}")
                    if TextLoader is not None:
                        loader = TextLoader(str(file_path), encoding="utf-8")
                        docs = loader.load()
                        for d in docs:
                            md = getattr(d, "metadata", {}) or {}
                            md.update({"source": document_source, "file": file_path.name})
                            d.metadata = md
                        all_docs.extend(docs)
                    else:
                        content = file_path.read_text(encoding="utf-8")
                        all_docs.append(Document(page_content=content, metadata={"source": document_source, "file": file_path.name}))
                    processed_files += 1
                else:
                    msg = f"不支持的文件格式: {file_path.suffix} ({file_path.name})"
                    logger.warning(msg)
                    errors.append(msg)
            except Exception as e:
                msg = f"加载文件失败 {file_path.name}: {e}"
                logger.error(msg)
                errors.append(msg)

        total_chunks = 0
        if all_docs:
            splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            chunks = splitter.split_documents(all_docs)
            for c in chunks:
                md = getattr(c, "metadata", {}) or {}
                md.setdefault("source", document_source)
                c.metadata = md

            self.vectorstore.add_documents(chunks)
            self._persist()
            self._update_bm25_with_chunks(chunks)
            total_chunks = len(chunks)
            logger.info(f"添加 {total_chunks} 个文本块到向量存储")

        return {"processed_files": processed_files, "total_chunks": total_chunks, "errors": errors}

    def retrieve(self, query: str, k: int = 5) -> List[Dict[str, Any]]:
        try:
            results = self.vectorstore.similarity_search_with_score(query, k=k)
            out = []
            for doc, score in results:
                out.append({"content": doc.page_content, "metadata": doc.metadata, "score": score})
            return out
        except Exception as e:
            logger.error(f"检索失败: {e}")
            return []

    def query(self, query: str, top_k: int = 5, score_threshold: float = 0.3,
              mode: str = 'hybrid') -> List[Dict[str, Any]]:
        try:
            if mode == 'bm25':
                results = self._bm25_search(query, k=top_k)
            elif mode == 'vector':
                results = self.vectorstore.similarity_search_with_score(query, k=top_k)
            else:
                results = self._hybrid_search(query, k=top_k)

            out = []
            for doc, score in results:
                if mode == 'hybrid':
                    sim = min(float(score) * RRF_K, 1.0)
                elif mode == 'bm25':
                    sim = 1.0 / (1.0 + math.exp(-float(score) / 10.0))
                else:
                    sim = 1.0 / (1.0 + float(score))

                if sim < score_threshold:
                    continue

                metadata = doc.metadata or {}
                source = metadata.get("file") or metadata.get("source") or "unknown"

                out.append({
                    "content": doc.page_content,
                    "metadata": doc.metadata,
                    "source": source,
                    "similarity_score": sim,
                })
            return out
        except Exception as e:
            logger.error(f"query 失败: {e}")
            return []

    def clear_db(self) -> Dict[str, Any]:
        try:
            self.chroma_client.delete_collection(self.collection_name)
            self.collection = self.chroma_client.get_or_create_collection(
                name=self.collection_name,
                metadata={"hnsw:space": "cosine"},
            )
            self.vectorstore = self._init_vectorstore()
            self._bm25_index = None
            self._bm25_texts = []
            self._bm25_metadatas = []
            pkl = Path(self._bm25_pickle_path)
            if pkl.exists():
                pkl.unlink()
            logger.info("数据库已清空")
            return {"status": "success", "message": "所有文档已清除"}
        except Exception as e:
            logger.error(f"clear_db 失败: {e}")
            return {"status": "error", "message": str(e)}

    def get_context_for_query(self, query: str, top_k: int = 5, mode: str = 'hybrid') -> str:
        results = self.query(query, top_k=top_k, score_threshold=0.1, mode=mode)
        if not results:
            return ""
        parts = []
        for i, doc in enumerate(results, 1):
            parts.append(f"[{i}] {doc['content']}")
        return "\n\n".join(parts)

    def get_db_stats(self) -> Dict[str, Any]:
        try:
            count = self.collection.count()
            return {
                "collection_name": self.collection_name,
                "total_chunks": count,
                "db_path": self.vector_db_path,
                "bm25_indexed": len(self._bm25_texts),
                "bm25_available": self._bm25_index is not None,
            }
        except Exception as e:
            logger.error(f"获取统计信息失败: {e}")
            return {
                "collection_name": self.collection_name,
                "total_chunks": 0,
                "db_path": self.vector_db_path,
                "bm25_indexed": 0,
                "bm25_available": False,
            }


# 单例
_rag_manager: Optional[RAGManager] = None

def get_rag_manager(vector_db_path: str = "./vector_db") -> RAGManager:
    global _rag_manager
    if _rag_manager is None:
        _rag_manager = RAGManager(vector_db_path=vector_db_path)
    return _rag_manager

def reset_rag_manager():
    global _rag_manager
    _rag_manager = None
