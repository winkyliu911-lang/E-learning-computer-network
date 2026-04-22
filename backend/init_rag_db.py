"""
RAG 数据库初始化脚本
用于将内置的计算机网络课程文档导入向量数据库
"""

import os
import sys
from pathlib import Path

# 添加后端路径
sys.path.insert(0, os.path.dirname(__file__))

from rag_manager import get_rag_manager
import logging

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def init_builtin_documents():
    """
    初始化 RAG 数据库
    """
    logger.info("开始初始化 RAG 数据库...")

    try:
        rag_manager = get_rag_manager(vector_db_path="./vector_db")
        logger.info("✅ RAG Manager 初始化完成")
        
        # 获取 backend/documents 目录
        documents_dir = Path(__file__).parent / "documents"
        documents_dir.mkdir(exist_ok=True)
        
        # 创建必要的目录
        (Path(__file__).parent / "uploads").mkdir(exist_ok=True)
        
        logger.info(f"\n📂 文档目录: {documents_dir.absolute()}")
        logger.info("📝 请将您的 PDF、Word 或 Markdown 文件放在这个目录中")
        logger.info("   然后重新运行此脚本以加载文件")
        
        # 获取现有文件
        file_paths = (
            list(documents_dir.glob("*.pdf")) + 
            list(documents_dir.glob("*.docx")) + 
            list(documents_dir.glob("*.md"))
        )
        
        if file_paths:
            logger.info(f"\n📄 找到 {len(file_paths)} 个文档:")
            for fp in file_paths:
                logger.info(f"   - {fp.name}")
            
            # 尝试加载文档
            try:
                results = rag_manager.add_documents(
                    [str(p) for p in file_paths],
                    document_source="builtin"
                )
                logger.info(f"\n✅ 文档处理完成!")
                logger.info(f"   - 处理文件数: {results['processed_files']}")
                logger.info(f"   - 总块数: {results['total_chunks']}")
                if results['errors']:
                    for error in results['errors']:
                        logger.error(f"   - {error}")
            except Exception as e:
                logger.warning(f"⚠️ 文档加载遇到问题: {e}")
                logger.info("✅ 但 RAG 数据库已初始化，可以直接使用")
        else:
            logger.info("\n✅ RAG 数据库已就绪，等待文件上传")
        
        # 显示统计
        stats = rag_manager.get_db_stats()
        logger.info(f"\n📊 数据库统计:")
        logger.info(f"   - 集合名称: {stats['collection_name']}")
        logger.info(f"   - 总块数: {stats['total_chunks']}")
        
    except Exception as e:
        logger.error(f"❌ 初始化失败: {e}")
        raise


def create_sample_document(documents_dir: Path) -> Path:
    """创建示例 Word 文档用于演示"""
    try:
        from docx import Document

        # 创建示例 Word 文档
        doc = Document()
        doc.add_heading('计算机网络基础知识', 0)

        doc.add_heading('1. 什么是计算机网络?', level=1)
        doc.add_paragraph(
            '计算机网络是指将地理位置不同的具有独立功能的多台计算机及其外部设备，'
            '通过通信线路连接起来，在网络操作系统、网络管理软件及通信协议的管理和协调下，'
            '实现资源共享和信息传递的计算机系统。'
        )

        doc.add_heading('2. 网络的基本功能', level=1)
        doc.add_paragraph('• 资源共享：硬件资源和软件资源的共享')
        doc.add_paragraph('• 信息传递：实时通信和信息交换')
        doc.add_paragraph('• 集中管理：对网络中的数据和资源进行集中管理')
        doc.add_paragraph('• 负载均衡：分散处理任务，提高系统效率')

        doc.add_heading('3. 网络体系结构', level=1)
        doc.add_paragraph('常见的网络体系结构包括：')
        doc.add_paragraph('• OSI 参考模型：7 层结构')
        doc.add_paragraph('• TCP/IP 模型：4 层结构')

        doc.add_heading('4. 常见协议', level=1)
        doc.add_paragraph('• IP 协议：互联网协议')
        doc.add_paragraph('• TCP 协议：传输控制协议')
        doc.add_paragraph('• UDP 协议：用户数据包协议')
        doc.add_paragraph('• HTTP/HTTPS 协议：超文本传输协议')
        doc.add_paragraph('• DNS 协议：域名系统')

        sample_path = documents_dir / "sample_networking.docx"
        doc.save(str(sample_path))
        logger.info(f"✅ 示例文档已创建: {sample_path.name}")
        return sample_path

    except Exception as e:
        logger.error(f"❌ 创建示例文档失败: {e}")
        raise


if __name__ == "__main__":
    try:
        init_builtin_documents()
        logger.info("\n🎉 RAG 数据库初始化成功!")
    except Exception as e:
        logger.error(f"\n❌ 初始化失败: {e}")
        sys.exit(1)
