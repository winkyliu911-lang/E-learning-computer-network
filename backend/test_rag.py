#!/usr/bin/env python3
"""
RAG 功能测试脚本
演示完整的 RAG 工作流程
"""

import os
import sys
from pathlib import Path

# 添加后端路径
sys.path.insert(0, os.path.dirname(__file__))

from rag_manager import get_rag_manager
import logging

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def test_rag_workflow():
    """测试完整的 RAG 工作流程"""

    logger.info("=" * 60)
    logger.info("🧪 RAG 系统测试")
    logger.info("=" * 60)

    # 初始化 RAG 管理器
    logger.info("\n1️⃣  初始化 RAG 管理器...")
    try:
        rag_manager = get_rag_manager(vector_db_path="./vector_db_test")
        logger.info("✅ RAG 管理器初始化成功")
    except Exception as e:
        logger.error(f"❌ 初始化失败: {e}")
        return

    # 检查向量数据库状态
    logger.info("\n2️⃣  检查向量数据库状态...")
    stats = rag_manager.get_db_stats()
    logger.info(f"   - 集合: {stats['collection_name']}")
    logger.info(f"   - 块数: {stats['total_chunks']}")
    logger.info(f"   - 路径: {stats['db_path']}")

    # 如果数据库为空，创建示例文档
    if stats['total_chunks'] == 0:
        logger.info("\n3️⃣  数据库为空，创建示例文档...")

        from docx import Document

        # 创建示例文档
        doc = Document()
        doc.add_heading('计算机网络知识库', 0)

        doc.add_heading('TCP 三次握手', level=1)
        doc.add_paragraph(
            'TCP 三次握手是建立 TCP 连接的过程。'
            'SYN: 客户端发送同步信号，'
            'SYN-ACK: 服务器回复同步确认，'
            'ACK: 客户端发送确认信号。'
        )

        doc.add_heading('IP 地址', level=1)
        doc.add_paragraph(
            'IPv4 地址是 32 位的，由 4 个字节组成。'
            '格式为：xxx.xxx.xxx.xxx，范围是 0.0.0.0 到 255.255.255.255。'
            'IPv6 地址是 128 位的。'
        )

        doc.add_heading('DNS 协议', level=1)
        doc.add_paragraph(
            'DNS 是域名系统，用于将域名转换为 IP 地址。'
            'DNS 查询过程包括递归查询和迭代查询。'
        )

        sample_path = "./sample_network.docx"
        doc.save(sample_path)
        logger.info(f"✅ 示例文档已创建: {sample_path}")

        # 添加到向量数据库
        logger.info("\n4️⃣  添加文档到向量数据库...")
        results = rag_manager.add_documents([sample_path], document_source="test")
        logger.info(f"   - 处理文件数: {results['processed_files']}")
        logger.info(f"   - 总块数: {results['total_chunks']}")
        if results['errors']:
            logger.error(f"   - 错误: {results['errors']}")
    else:
        logger.info("\n3️⃣  ⏭️  跳过文档创建（数据库非空）")
        logger.info("\n4️⃣  ⏭️  跳过文档添加（已有数据）")

    # 测试查询
    logger.info("\n5️⃣  测试查询...")
    test_queries = [
        "TCP 三次握手是什么?",
        "IP 地址的格式是什么?",
        "DNS 的作用是什么?"
    ]

    for query in test_queries:
        logger.info(f"\n   📝 查询: {query}")
        results = rag_manager.query(query, top_k=3, score_threshold=0.2)

        if results:
            for i, doc in enumerate(results, 1):
                logger.info(
                    f"   [{i}] {doc['source']} (相似度: {doc['similarity_score']:.2%})"
                )
                logger.info(f"       内容: {doc['content'][:100]}...")
        else:
            logger.info("   ℹ️  未找到相关文档")

    # 获取上下文
    logger.info("\n6️⃣  获取查询上下文...")
    query = "TCP 握手"
    context = rag_manager.get_context_for_query(query)
    if context:
        logger.info(f"   ✅ 上下文长度: {len(context)} 字符")
        logger.info(f"   📄 上下文预览:\n{context[:300]}...")
    else:
        logger.info("   ℹ️  无相关上下文")

    # 统计
    logger.info("\n7️⃣  最终统计...")
    final_stats = rag_manager.get_db_stats()
    logger.info(f"   - 总块数: {final_stats['total_chunks']}")
    logger.info(f"   - 数据库路径: {final_stats['db_path']}")

    # 清理
    logger.info("\n8️⃣  清理测试文件...")
    test_file = Path("./sample_network.docx")
    if test_file.exists():
        test_file.unlink()
        logger.info("✅ 测试文件已删除")

    logger.info("\n" + "=" * 60)
    logger.info("✅ 测试完成！")
    logger.info("=" * 60)


def test_rag_with_custom_files():
    """测试使用自定义文件的 RAG"""
    logger.info("\n" + "=" * 60)
    logger.info("🧪 自定义文件测试")
    logger.info("=" * 60)

    documents_dir = Path("./documents")
    documents_dir.mkdir(exist_ok=True)

    logger.info(f"\n📂 检查文档目录: {documents_dir.absolute()}")

    file_paths = list(documents_dir.glob("*.pdf")) + list(documents_dir.glob("*.docx"))

    if not file_paths:
        logger.info("⚠️  未找到 PDF 或 DOCX 文件")
        logger.info(f"   请在 {documents_dir.absolute()} 中放置文件")
        return

    logger.info(f"✅ 找到 {len(file_paths)} 个文件:")
    for fp in file_paths:
        logger.info(f"   - {fp.name}")

    # 处理文件
    logger.info("\n🔄 处理文件...")
    rag_manager = get_rag_manager()
    results = rag_manager.add_documents([str(p) for p in file_paths])

    logger.info(f"✅ 处理完成!")
    logger.info(f"   - 处理文件数: {results['processed_files']}")
    logger.info(f"   - 总块数: {results['total_chunks']}")

    if results['errors']:
        logger.error(f"   - 错误数: {len(results['errors'])}")
        for error in results['errors']:
            logger.error(f"     • {error}")


if __name__ == "__main__":
    # 运行示例测试
    test_rag_workflow()

    # 如果有自定义文件，运行自定义测试
    if Path("./documents").exists():
        files = list(Path("./documents").glob("*.pdf")) + list(Path("./documents").glob("*.docx"))
        if files:
            test_rag_with_custom_files()

    logger.info("\n💡 提示:")
    logger.info("1. 要使用 RAG，请在 documents/ 目录中放置 PDF 或 DOCX 文件")
    logger.info("2. 运行 init_rag_db.py 初始化数据库")
    logger.info("3. 启动 Flask 后端: python3 app.py")
    logger.info("4. 在前端启用 RAG 选项进行聊天")
